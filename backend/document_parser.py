import os
import re
from typing import Dict, Any
from docx import Document as DocxDocument
from PyPDF2 import PdfReader
import logging

logger = logging.getLogger(__name__)

class DocumentParser:
    """Parse different document types and extract text content"""
    
    @staticmethod
    def extract_text(file_path: str) -> Dict[str, Any]:
        """
        Extract text from various file formats
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        try:
            file_extension = os.path.splitext(file_path)[1].lower()
            
            if file_extension == '.pdf':
                return DocumentParser._extract_from_pdf(file_path)
            elif file_extension in ['.docx', '.doc']:
                return DocumentParser._extract_from_docx(file_path)
            elif file_extension == '.txt':
                return DocumentParser._extract_from_txt(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return {
                'text': '',
                'page_count': 0,
                'word_count': 0,
                'error': str(e)
            }
    
    @staticmethod
    def _extract_from_pdf(file_path: str) -> Dict[str, Any]:
        """Extract text from PDF file"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                text = ""
                
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                
                return {
                    'text': text.strip(),
                    'page_count': len(pdf_reader.pages),
                    'word_count': len(text.split()),
                    'format': 'pdf'
                }
        except Exception as e:
            raise Exception(f"Failed to extract PDF text: {str(e)}")
    
    @staticmethod
    def _extract_from_docx(file_path: str) -> Dict[str, Any]:
        """Extract text from Word document"""
        try:
            doc = DocxDocument(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                text += "\n"
            
            return {
                'text': text.strip(),
                'page_count': 1,  # Word docs don't have clear page breaks in python-docx
                'word_count': len(text.split()),
                'format': 'docx'
            }
        except Exception as e:
            raise Exception(f"Failed to extract Word document text: {str(e)}")
    
    @staticmethod
    def _extract_from_txt(file_path: str) -> Dict[str, Any]:
        """Extract text from plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            return {
                'text': text.strip(),
                'page_count': 1,
                'word_count': len(text.split()),
                'format': 'txt'
            }
        except Exception as e:
            raise Exception(f"Failed to extract text file: {str(e)}")
    
    @staticmethod
    def clean_text(text: str) -> str:
        """Clean and normalize text for analysis"""
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep punctuation
        text = re.sub(r'[^\w\s\.\,\!\?\;\:\-\(\)\"\'%\$]', '', text)
        
        return text.strip()
    
    @staticmethod
    def get_text_statistics(text: str) -> Dict[str, Any]:
        """Get basic statistics about the text"""
        words = text.split()
        sentences = re.split(r'[.!?]+', text)
        
        return {
            'character_count': len(text),
            'word_count': len(words),
            'sentence_count': len([s for s in sentences if s.strip()]),
            'average_word_length': sum(len(word) for word in words) / len(words) if words else 0,
            'average_sentence_length': len(words) / len(sentences) if sentences else 0
        } 